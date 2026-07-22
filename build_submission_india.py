# -*- coding: utf-8 -*-
import sys
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()

style = d.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

section = d.sections[0]
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def h1(text, size=14, space_before=6, space_after=4):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def h2(text, size=11.5):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def body(text, size=9.5, space_after=4):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def bullet(text, size=9.5):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


# ================= TITLE BLOCK =================
title = d.add_paragraph()
title.paragraph_format.space_after = Pt(0)
r = title.add_run("Predicting Churn Among High-Value Prepaid Subscribers")
r.bold = True
r.font.size = Pt(16)

sub = d.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("Final Submission: Working Prototype, Testing and Recommendation (Section 4.2 / Section 5) -- Indian Prepaid Telecom Market")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = d.add_paragraph()
meta.paragraph_format.space_after = Pt(8)
r = meta.add_run("Group C8: Shikha Arora, Harsh Arora, Kaushik Suman, Shubham Jaiswal  |  AI & Analytics for Managers")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ================= THE PROCESS =================
h1("The Process")

h2("1. Process and its value")
body(
    "We score churn risk for high-value prepaid subscribers of an Indian telecom operator. "
    "Input: three months of a customer's recharge and usage behaviour, average revenue per user, "
    "recharge amount, outgoing and incoming call minutes, 2G/3G data volume, and age on network. "
    "Output: a churn risk score plus a suggested retention action. This matters more in India's "
    "prepaid market than it would for a postpaid contract, since a prepaid customer can leave with "
    "zero notice by simply not recharging, and by the time a number goes silent it is usually too "
    "late to win them back. We also restrict scoring to high-value customers, the top 30 percent by "
    "recharge spend, since that group accounts for roughly 80 percent of revenue in the Indian and "
    "Southeast Asian market, so protecting them protects most of the revenue at risk. Success is "
    "judged on precision-recall performance and top-20-percent capture rate rather than plain "
    "accuracy, since only 8.6 percent of high-value customers actually churn."
)

h2("2. Current state and positioning")
body(
    "Prepaid churn in India runs at real scale. TRAI's own subscription reports show mobile number "
    "portability requests sustained at 14 to 16 million a month since mid-2025, and Vodafone Idea "
    "lost about 2.2 million active subscribers in a single month as tariff hikes pushed customers "
    "toward BSNL and Airtel. Because prepaid has no cancellation event to trigger on, the "
    "traditional response has been reactive and broad: bulk recharge-reminder SMS blasts and generic "
    "win-back offers sent to any number that goes quiet for a set number of days, the same offer to "
    "everyone regardless of why they actually stopped recharging."
)
body(
    "Indian operators are already moving past this. Airtel's AI-led customer-experience overhaul "
    "with Amdocs cut call-centre volumes 60 percent and order fallout 90 percent in six months, and "
    "Airtel's own Xtelify subsidiary now sells an AI platform marketed explicitly on reducing churn "
    "and lifting ARPU. Reliance Jio built JioBrain in-house over two years, running customer "
    "experience analytics across roughly 500 million subscribers. Our prototype replicates the core "
    "idea these platforms sell, a live per-customer risk score from usage and recharge behaviour, but "
    "built and shown end to end ourselves rather than a vendor result taken on faith: a real "
    "published dataset, six benchmarked configurations compared honestly, and a live API. We extend "
    "beyond typical classroom versions of this problem by treating the 8.6 percent class imbalance "
    "and the leakage risk as first-class design decisions, not afterthoughts."
)

# ================= THE PROTOTYPE =================
h1("The Prototype")

h2("3. Approach")
body(
    "Three deliberate design choices were baked in before any model got fit. First, churn is "
    "usage-based, not contract-based: a high-value customer counts as churned if calls and data "
    "usage both drop to zero in the target month, since prepaid has no notice period to define churn "
    "any other way. Second, only months 6 to 8 of behaviour are used as features to predict month 9; "
    "every column ending in _9 is dropped from the model inputs entirely, since those columns exist "
    "only to build the label and leaving them in would let the model see the answer first. Third, "
    "with only 8.6 percent of high-value customers churning, every model was trained with class "
    "weighting (scale_pos_weight for XGBoost, class_weight=\"balanced\" for the others) so it cannot "
    "just predict \"stays\" for everyone and still look accurate. We benchmarked three model types on "
    "two feature sets each and picked the deployed configuration on PR-AUC and top-20-percent capture "
    "rather than raw ROC-AUC, for reasons Section 5 makes concrete. That configuration is saved as a "
    "real scikit-learn pipeline and served live through a FastAPI backend, so the prototype calls the "
    "actual trained model over a REST endpoint. Retention-offer matching stays a rule-based layer, "
    "since no public dataset records which Indian recharge offer was given to whom and whether it worked."
)

h2("4. Input data")
body(
    "The \"Telecom Churn Case Study\" dataset: 99,999 real prepaid customers of an unnamed Indian "
    "operator, a widely published case study used across Indian data-science coursework (independently "
    "verifiable: several public solutions to this exact dataset are mirrored on GitHub). We filtered "
    "to high-value customers, average recharge across months 6 and 7 at or above the 70th percentile "
    "(Rs.368.50), leaving 30,011 customers, 8.64 percent of whom churned by the documented usage-based "
    "definition. Fields are real call-detail-record-style metrics in INR: recharge amount, ARPU, "
    "on-net/off-net/local/STD/roaming minutes, 2G/3G data volume, and age on network, tracked monthly. "
    "About 75 percent of the data-recharge columns are missing, which the documentation defines as "
    "\"no data recharge that month,\" not a quality problem, so we imputed those as zero. We also "
    "engineered four decline features (months 6-7 average minus month 8, for ARPU, recharge, and both "
    "call directions), matching the case study's own concept of an \"action phase\" before churn."
)

h2("5. Output and accuracy")
body(
    "The prototype outputs a churn probability, a risk tier, and a suggested action. Our selected "
    "configuration, XGBoost on the full 143-feature set, reached 0.937 ROC-AUC, 0.697 PR-AUC, 90.5 "
    "percent accuracy, 47.2 percent precision, and 82.4 percent recall on held-out test data. We did "
    "not simply take the ROC-AUC leader: Random Forest scored marginally higher on ROC-AUC (0.9376 "
    "versus 0.9369), but with churn this rare, ROC-AUC is dominated by how well a model ranks the "
    "large majority class and can look similar across models that differ meaningfully on the minority "
    "class that actually matters. XGBoost wins on PR-AUC (0.697 versus 0.689) and on the number that "
    "maps directly to a retention team's capacity: calling the top 20 percent of high-value customers "
    "by risk score catches 87 percent of everyone who actually churns."
)

h2("6. Models tested")
body(
    "Every combination of three model types (Logistic Regression, Random Forest, XGBoost, all "
    "class-weighted for the 8.6 percent churn rate) against two feature sets (a 19-field core set: "
    "ARPU, recharge, minutes of use, data volume, tenure; and the full 143-field set with granular "
    "call-type breakdowns and decline features), six configurations, identical train/test split."
)

table_rows = [
    ("Feature set", "Model", "Features", "ROC-AUC", "PR-AUC", "Recall", "Top-20% capture"),
    ("Full fields", "XGBoost", "143", "0.937", "0.697", "82.4%", "87.0%"),
    ("Full fields", "Random Forest", "143", "0.938", "0.689", "79.0%", "86.1%"),
    ("Core fields", "XGBoost", "19", "0.914", "0.658", "77.0%", "82.1%"),
    ("Core fields", "Random Forest", "19", "0.911", "0.652", "70.7%", "80.7%"),
    ("Core fields", "Logistic Reg.", "19", "0.886", "0.461", "84.4%", "80.4%"),
    ("Full fields", "Logistic Reg.", "143", "0.892", "0.461", "83.5%", "80.1%"),
]
tbl = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = True
for ri, row in enumerate(table_rows):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(val)
        run.font.size = Pt(9)
        if ri == 0:
            run.bold = True
            set_cell_shading(cell, "1F5C4A")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif ri == 1:
            set_cell_shading(cell, "E1F5EE")
            run.bold = True
d.add_paragraph().paragraph_format.space_after = Pt(2)
cap = d.add_paragraph()
cap.paragraph_format.space_after = Pt(6)
r = cap.add_run("Table 1. All six model x feature-set combinations on held-out test data, ranked by top-20% capture rate (the metric we selected on). Deployed configuration highlighted.")
r.italic = True
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

h2("7. Prototype")
body(
    "Screenshots below are of the actual working tool, not a mock-up: a browser front end that "
    "sends a customer's three-month recharge and usage profile (entered as ARPU, recharge amount, "
    "outgoing/incoming minutes for months 6 to 8, plus latest data usage, roaming, and tenure) to a "
    "FastAPI backend (api.py), which runs it through the real fitted pipeline and returns a live "
    "prediction, shown in the \"Live model: connected\" badge. The ranked list of 24 real held-out "
    "high-value customers (mobile numbers masked) is scored by that same model against their actual "
    "outcome. The whole service, API and prototype page together, is packaged with a Dockerfile so it "
    "can be deployed to a public host rather than only run locally; the HTML detects whether it is "
    "opened as a local file or served by the API itself and calls the right address either way."
)

img_table = d.add_table(rows=1, cols=2)
img_table.autofit = True
c0 = img_table.cell(0, 0)
c0.paragraphs[0].add_run().add_picture("screenshots/calculator_final.png", width=Inches(3.15))
c1 = img_table.cell(0, 1)
c1.paragraphs[0].add_run().add_picture("screenshots/crop_rankedlist.png", width=Inches(3.15))
cap2 = d.add_paragraph()
cap2.paragraph_format.space_before = Pt(2)
cap2.paragraph_format.space_after = Pt(6)
r = cap2.add_run("Left: live what-if calculator (month-6/7/8 recharge and usage grid). Right: 24 real held-out high-value customers ranked by live risk score, with actual outcomes.")
r.italic = True
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

h2("8. Tools used")
body(
    "Claude Code to write and run the full pipeline, the API, and the browser prototype. Python with "
    "pandas and scikit-learn for data prep, feature engineering, and the Logistic Regression and "
    "Random Forest models, and xgboost for the selected model. Matplotlib for the benchmark charts. "
    "The deployed pipeline is saved with joblib and served through FastAPI and uvicorn, with the same "
    "service also hosting the prototype page; the front end is plain HTML, CSS, and JavaScript "
    "calling that API with fetch(). Docker packages the whole service for hosting on a platform such "
    "as Render or Railway. Screenshots were captured with headless Google Chrome while the API ran."
)

h2("9. Benchmark comparison")
body(
    "Because this is a well-known, independently published case study, we can compare against other "
    "public solutions to this exact dataset rather than a generic industry baseline. A commonly cited "
    "public solution reports 80 percent recall at 86 percent accuracy, 37 percent precision, and an "
    "F1 of 0.51, using Random Forest after PCA. Our selected XGBoost configuration reaches 82.4 "
    "percent recall at a meaningfully better 47.2 percent precision and 0.60 F1, on the same "
    "underlying data and churn definition, likely from the added decline features and class-weight "
    "tuning. What we fully replicate and test: the case study's own classification task, high-value "
    "churn from months 6-8 predicting month 9, end to end, on real held-out data. What we do not "
    "replicate: an uplift or offer-response model, since no public dataset records which Indian "
    "recharge offer was given to whom and whether it worked; that stays a reasoned, rule-based design "
    "choice rather than a tested model."
)

# ================= REFLECTION =================
h1("Reflection")

h2("10. Key challenges")
bullet(
    "Defining churn honestly for prepaid: there is no cancellation event to label against, so we "
    "used the dataset's own usage-based definition (zero calls and zero data in the target month) "
    "rather than inventing a friendlier one."
)
bullet(
    "Avoiding leakage: the target month's own columns would trivially predict that month's churn "
    "label, so every column ending in _9 was dropped from the model inputs entirely, not just the "
    "columns used to build the label."
)
bullet(
    "Class imbalance: with only 8.6 percent of high-value customers churning, an unweighted model "
    "could reach 91 percent accuracy by predicting \"stays\" for everyone; class weighting and "
    "choosing PR-AUC/capture rate over raw ROC-AUC as the selection metric were both necessary, not "
    "optional extras."
)
bullet(
    "Making a 143-feature model usable in a live demo: a real customer profile has far more fields "
    "than anyone would hand-type into a form, so the API defaults non-exposed fields to their "
    "dataset median and only exposes about 16 headline behavioural signals live."
)

h2("11. Key learnings")
bullet(
    "The metric you optimise for changes the winner: XGBoost narrowly loses to Random Forest on "
    "ROC-AUC but wins on PR-AUC and top-20-percent capture, the two numbers that actually matter "
    "under 8.6 percent churn, so we deployed XGBoost, not the ROC-AUC leaderboard topper."
)
bullet(
    "Feature engineering beat raw features: the four decline features (recent behaviour against the "
    "earlier average) showed up among the model's most important signals, confirming the case "
    "study's own idea of an \"action phase\" before churn is real, not just a theoretical framing."
)
bullet(
    "Our numbers beat a commonly cited public solution to the same dataset (82.4 percent recall and "
    "47.2 percent precision versus a widely referenced 80 percent recall and 37 percent precision), "
    "which is a genuinely checkable claim since both solve the identical published task."
)
bullet(
    "A live, deployable prototype surfaces different problems than a notebook does: getting a "
    "143-feature model to work behind a usable form, and packaging it so it can run on a public host "
    "rather than only our own laptop, took real design decisions a static report would never force."
)

out_name = sys.argv[1] if len(sys.argv) > 1 else "Churn Risk Scoring - Final Submission.docx"
d.save(out_name)
print(f"Saved: {out_name}")
